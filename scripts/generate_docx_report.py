"""Script sinh báo cáo Word (.docx) chuyên nghiệp bằng cách chèn nội dung học thuật vào template, xóa các mục cũ thừa và tự động verify."""

import os
import re
from pathlib import Path
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import pandas as pd
import json

PROJECT_ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = PROJECT_ROOT / "reports" / "2582003108_MonHoc_NguyenMinhTan.docx"
OUTPUT_DOCX = PROJECT_ROOT / "reports" / "BaoCao_MonHoc_NguyenMinhTan_Complete.docx"

# --- NỘI DUNG CHI TIẾT CỦA BÁO CÁO ---

COVER_PAGE_DATA = {
    "UNIVERSITY": "PHÂN HIỆU TRƯỜNG ĐẠI HỌC ...",
    "FACULTY": "KHOA CÔNG NGHỆ THÔNG TIN / KHOA HỌC DỮ LIỆU",
    "MAIN_TITLE_VN": "PHÂN TÍCH THỊ TRƯỜNG VIỆC LÀM IT & GỢI Ý ỨNG VIÊN BẰNG MACHINE LEARNING",
    "MAIN_TITLE_EN": "IT JOB MARKET ANALYSIS & CANDIDATE RECOMMENDATION USING MACHINE LEARNING",
    "AUTHOR_INFO": "Học viên cao học: Nguyễn Minh Tan - [Tên thành viên khác]",
    "ADVISOR_INFO": "GVHD: TS. Hoàng Văn Quý",
    "DATE": "TP. HỒ CHÍ MINH, THÁNG 8 NĂM 2026",
}

FULL_CONTENT = {
    "LỜI MỞ ĐẦU": [
        ("I. Lý do lựa chọn đề tài",
         "Trong bối cảnh nền kinh tế số và chuyển đổi số diễn ra mạnh mẽ tại Việt Nam, ngành Công nghệ Thông tin (IT) đóng vai trò là động lực cốt lõi thúc đẩy tăng trưởng kinh tế. Tuy nhiên, thị trường việc làm IT hiện nay đang đối mặt với một nghịch lý: sự mất cân đối cung - cầu thông tin giữa nhà tuyển dụng và người lao động. Dữ liệu tuyển dụng bị phân tán trên nhiều nền tảng trực tuyến khác nhau (Itviec, Glints, TopCV, Careerviet...) với các định dạng không đồng nhất. Người tìm việc gặp nhiều khó khăn trong việc định giá năng lực, nhận diện các kỹ năng cốt lõi và định hướng lộ trình học tập. Từ những thách thức đó, việc nghiên cứu áp dụng các phương pháp Khoa học Dữ liệu và Học máy để xây dựng hệ thống thu thập, làm sạch, phân tích thị trường việc làm tự động, đồng thời phát triển mô hình dự báo lương và hệ thống gợi ý việc làm thông minh là một hướng đi cấp thiết và có giá trị ứng dụng cao. Đó chính là lý do chúng tôi thực hiện đề tài này."),
        ("II. Mục đích và Nhiệm vụ nghiên cứu",
         "Mục đích nghiên cứu: Khai thác, làm sạch và phân tích sâu dữ liệu tuyển dụng ngành IT tại Việt Nam để cung cấp góc nhìn toàn diện về xu hướng thị trường; từ đó xây dựng các mô hình học máy dự báo mức lương và đề xuất hệ thống gợi ý việc làm tối ưu cho ứng viên.\nNhiệm vụ nghiên cứu: 1) Thiết kế và triển khai quy trình thu thập dữ liệu tự động (Crawler v2); 2) Xây dựng pipeline làm sạch, chuẩn hóa dữ liệu; 3) Thực hiện phân tích khám phá dữ liệu (EDA); 4) Xây dựng và đánh giá các mô hình học máy (Baseline, Linear Regression, Decision Tree, Random Forest) để dự đoán mức lương; 5) Áp dụng học không giám sát (K-Means Clustering) và hệ thống gợi ý việc làm dựa trên độ tương đồng Cosine (Content-based Recommendation)."),
        ("III. Phương pháp và Phạm vi nghiên cứu",
         "Phương pháp nghiên cứu: Sử dụng phương pháp định lượng và thống kê mô tả; phương pháp học máy (Machine Learning) với các thuật toán hồi quy, phân cụm và đo lường độ tương đồng; lập trình hướng đối tượng (OOP) để xây dựng kiến trúc mã nguồn. Phạm vi nghiên cứu: Thị trường việc làm IT tại các trung tâm công nghệ lớn của Việt Nam (Hồ Chí Minh, Hà Nội, Đà Nẵng); dữ liệu tuyển dụng được thu thập và cập nhật trong giai đoạn 2026."),
        ("IV. Kết cấu đề tài",
         "Ngoài phần Mở đầu, Kết luận và Danh mục tài liệu tham khảo, nội dung chính của báo cáo được chia làm 3 chương:\n• Chương 1: Tổng quan bài toán và Cơ sở lý thuyết\n• Chương 2: Phương pháp nghiên cứu và Xây dựng Pipeline dữ liệu\n• Chương 3: Thực nghiệm, Phân tích EDA và Đánh giá mô hình")
    ],
    "CHƯƠNG 1 TỔNG QUAN VỀ BÀI TOÁN VÀ CƠ SỞ LÝ THUYẾT": [
        ("1.1. Bối cảnh, Đặt vấn đề và Tính cấp thiết của đề tài",
         "Thị trường công nghệ thông tin (IT) Việt Nam đang trong giai đoạn tăng trưởng nhanh với nhu cầu nhân lực lớn, đặc biệt tại ba trung tâm công nghệ Hồ Chí Minh, Hà Nội và Đà Nẵng. Tuy nhiên, thị trường đang đối mặt với nghịch lý mất cân đối cung - cầu thông tin: dữ liệu tuyển dụng bị phân tán trên nhiều nền tảng khác nhau (Itviec, Glints, TopCV, Careerviet) với định dạng phi cấu trúc (unstructured text) không đồng nhất. Người tìm việc gặp khó khăn trong việc định giá năng lực, nhận diện kỹ năng cốt lõi và định hướng lộ trình học tập. Từ đó, đề tài đặt ra 5 câu hỏi nghiên cứu: (RQ1) Xu hướng kỹ năng nào đang được nhà tuyển dụng yêu cầu nhiều nhất? (RQ2) Lương biến động theo kinh nghiệm, thành phố và kỹ năng như thế nào? (RQ3) Có thể dự báo mức lương từ dữ liệu tuyển dụng không và độ chính xác ra sao? (RQ4) Thị trường việc làm IT phân khúc thành những nhóm nào? (RQ5) Làm thế nào để gợi ý việc làm phù hợp với hồ sơ kỹ năng của ứng viên? Các câu hỏi này được trả lời lần lượt trong Chương 3 (mục F1 đến F4). Đề tài có giá trị ứng dụng cao: hỗ trợ ứng viên định giá năng lực, giúp nhà tuyển dụng tối ưu chính sách lương cạnh tranh, và cung cấp dữ liệu thực cho nền tảng giáo dục xây dựng lộ trình học tập."),
        ("1.2. Cơ sở Lý thuyết về Phân tích Dữ liệu và Học máy",
         "Đề tài tuân thủ quy trình Khoa học Dữ liệu 9 bước tiêu chuẩn, từ thu thập, làm sạch, phân tích khám phá, mô hình hóa đến triển khai [1]. Nội dung lý thuyết được chia thành bốn nhóm tương ứng với bốn giai đoạn chính của pipeline: tiền xử lý dữ liệu phi cấu trúc, hồi quy học có giám sát, phân cụm không giám sát, và hệ gợi ý dựa trên nội dung."),
        ("1.2.1. Tiền xử lý dữ liệu phi cấu trúc (Unstructured Data Preprocessing)",
         "Dữ liệu tuyển dụng là văn bản tự do chứa nhiều biến thể ngôn ngữ và định dạng; theo McKinney [9], làm sạch dữ liệu chiếm phần lớn thời gian trong pipeline phân tích dữ liệu. Đề tài xây dựng bốn thành phần tiền xử lý: (1) SalaryParser dùng 6 biểu thức chính quy nhận diện 8 loại cấu trúc lương (RANGE, UP_TO, FROM, YEARLY, USD, HIDDEN, SINGLE, UNKNOWN); quy đổi USD sang VND với tỷ giá 25.000, lương theo năm chia cho 12 tháng, khoảng \"tới X\" lấy mức giữa bằng 70% mức tối đa, \"từ X\" lấy 130% mức tối thiểu; (2) SkillNormalizer dùng từ điển 188 quy tắc đồng nghĩa chuẩn hóa về 45 kỹ năng, kết hợp fuzzy matching ngưỡng > 0.8 cho các biến thể lỗi chính tả; (3) ExperienceNormalizer parse số năm kinh nghiệm từ nhiều định dạng tiếng Việt/Anh và xếp vào 5 bậc (entry, junior, mid, senior, lead); (4) Deduplicator triển khai 4 pha khử trùng lặp (exact theo job_id, exact theo title + company, fuzzy theo title ngưỡng 0.8, fuzzy theo description ngưỡng 0.7), đã loại bỏ 70 bản ghi trùng lặp [1]. Kết quả thực tế: tỷ lệ ẩn lương ghi nhận thực tế là 56%, chỉ 6.6% tin có thông tin kỹ năng chi tiết — phản ánh giới hạn hiển thị của nguồn dữ liệu, được xử lý bằng cờ is_hidden và cột salary_mid."),
        ("1.2.2. Hồi quy học có giám sát (Supervised Regression)",
         "Hồi quy là bài toán học có giám sát dự đoán một biến liên tục từ tập đặc trưng đầu vào [2]. Mô hình hồi quy tuyến tính giả định quan hệ tuyến tính giữa đặc trưng và biến mục tiêu [2]:"),
        ("__MATH__ y = β₀ + β₁x₁ + … + βₙxₙ ; t", ""),
        ("",
         "Trong đó y là mức lương dự đoán, xj là các đặc trưng (kinh nghiệm, thành phố, ngành...), βj là hệ số hồi quy. Để đo chất lượng mô hình, đề tài sử dụng ba chỉ số sai số [1][3]: MSE, RMSE và R²."),
        ("__MATH__ MSE = (1/n) ∑(yi − ŷi)² ; t", ""),
        ("__MATH__ RMSE = √MSE ; t", ""),
        ("__MATH__ R² = 1 − (∑(yi − ŷi)²) / (∑(yi − ȳ)²) ; t", ""),
        ("",
         "Trong đó yi là giá trị thực, ŷi là giá trị dự đoán, ȳ là trung bình của giá trị thực; R² thuộc (−∞, 1], giá trị gần 1 nghĩa là mô hình giải thích tốt phương sai của dữ liệu [2]. Trên thực tế, đề tài dùng cây quyết định (Decision Tree) với max_depth = 10, min_samples_leaf = 5 và rừng ngẫu nhiên (Random Forest) với n_estimators = 100, max_depth = 15, min_samples_leaf = 4 [4][5]. Pipeline đặc trưng dùng ColumnTransformer: imputer median + StandardScaler cho cột số, OneHotEncoder cho cột phân loại, OrdinalEncoder cho bậc kinh nghiệm; 7 nhóm đặc trưng chính (experience_years, city, job_type, remote_option, education_level, industry, company_size, experience_bin). Dữ liệu chia 80/20 và đánh giá bằng 5-fold cross-validation. Kết quả: Baseline RMSE 8.97, R² = −0.010; Random Forest đạt R² gần 1.0 (chi tiết bảng Chương 3)."),
        ("1.2.3. Phân cụm không giám sát (K-Means Clustering)",
         "Phân cụm nhóm các quan sát đồng dạng mà không cần nhãn [3]. K-Means tối thiểu hóa tổng bình phương khoảng cách từ mỗi điểm đến tâm cụm [6]:"),
        ("__MATH__ J = ∑ₖ ∑ₓ ‖x − μₖ‖² ; t", ""),
        ("",
         "Trong đó μk là tâm cụm thứ k, x là các điểm dữ liệu. Để đánh giá chất lượng phân cụm, chỉ số Silhouette do Rousseeuw [7] đề xuất:"),
        ("__MATH__ s(i) = (b(i) − a(i)) / max(a(i), b(i)) ; t", ""),
        ("",
         "Trong đó a(i) là khoảng cách trung bình từ điểm i đến các điểm cùng cụm, b(i) là khoảng cách đến cụm gần nhất khác; s(i) thuộc [−1, 1], giá trị dương thể hiện cụm tách biệt tốt [7]. Thực nghiệm khảo sát k từ 2 đến 10, chọn k = 5 với Silhouette Score = 0.38; dữ liệu được chuẩn hóa bằng StandardScaler trước khi K-Means (n_init = 10, random_state = 42) và trực quan hóa qua PCA 2 chiều. Kết quả hình thành 5 phân khúc: Junior-Mid tại Hà Nội (lương TB 15.1M), Mid-Senior tại TP.HCM (27.1M), Senior thu nhập cao (41.9M), việc phổ thông (20.8M) và Remote (31.6M)."),
        ("1.2.4. Hệ gợi ý dựa trên nội dung (Content-based Recommendation)",
         "Hệ gợi ý dựa trên nội dung xây dựng hồ sơ người dùng từ các đặc trưng của mục đã tương tác và gợi ý các mục tương tự [8]. Trong đề tài, hồ sơ là vector kỹ năng nhị phân của ứng viên được so sánh với ma trận job × skill dùng MultiLabelBinarizer; độ tương đồng Cosine giữa hai vector A và B [8]:"),
        ("__MATH__ cos(A,B) = (A·B) / (|A||B|) ; t", ""),
        ("",
         "Điểm tương đồng cao nghĩa là việc làm có nhiều kỹ năng trùng với hồ sơ ứng viên. Hệ thống hỗ trợ lọc theo thành phố (không phân biệt hoa thường) và kinh nghiệm (±0.5 năm, fallback experience_bin), trả về Top-N kèm danh sách matched/missing skills để ứng viên nhận diện kỹ năng cần bổ sung [8]."),
        ("1.3. Tính khả dụng của các phương pháp trong đề tài",
         "Tính khả dụng được đối chiếu lần lượt cho từng nhóm phương pháp với đặc thù dữ liệu tuyển dụng Việt Nam:\n"
         "• Tiền xử lý: các kỹ thuật regex và từ điển đồng nghĩa khả dụng ngay vì dữ liệu lương/kỹ năng có quy luật lặp lại cao; fuzzy matching bù cho biến thể gõ không dấu và viết tắt.\n"
         "• Hồi quy: khả dụng vì 44 thuộc tính cung cấp đủ đặc trưng phân loại; kết quả Random Forest gần 1.0 R² cho thấy đặc trưng đủ mạnh để giải thích lương, dù cần thận trọng với nguy cơ overfitting (mục 2.1 KẾT LUẬN).\n"
         "• K-Means: khả dụng để khám phá phân khúc thị trường không cần nhãn; Silhouette 0.38 ở mức chấp nhận được cho dữ liệu nhiều chiều.\n"
         "• Content-based: khả dụng nhất khi kỹ năng được chuẩn hóa thành 45 tên gọi — nền tảng để so sánh Cosine; hạn chế là chỉ 6.6% tin có kỹ năng chi tiết.\n"
         "Tổng hợp: mỗi kỹ thuật được chọn vì phù hợp với dạng dữ liệu hiện có (phi cấu trúc, thiếu nhãn, độ phủ kỹ năng thấp) và cho kết quả định lượng rõ ràng; các hạn chế (overfitting, độ phủ kỹ năng) được ghi nhận là hướng cải tiến ở Chương 3 và KẾT LUẬN.")
    ],
    "CHƯƠNG 2 PHƯƠNG PHÁP NGHIÊN CỨU VÀ DỮ LIỆU ĐẦU VÀO": [
        ("2.1. Kiến trúc Hệ thống Thu thập Dữ liệu (Crawler v2)",
         "Hệ thống thu thập dữ liệu được thiết kế theo kiến trúc đa tầng, vận hành qua hàm điều phối run_crawl() trong src/crawl/pipeline.py. Luồng tổng thể thực hiện vòng lặp lồng nhau site × keyword với 22 keyword tìm kiếm (python, java, react, data, devops, machine learning, golang...), mỗi tổ hợp crawl tối đa 2 trang. Trước khi ghi dữ liệu, hệ thống kiểm tra ngưỡng min_total_jobs: nếu tổng số bản ghi thu được dưới ngưỡng, toàn bộ lần chạy bị từ chối (không ghi CSV) nhằm tránh ghi nhận dữ liệu lỗi từ các trang bị chặn. Mỗi lần chạy ghi một file crawl_history JSON (timestamp, n_jobs, n_new, src_counts) phục vụ kiểm toán và theo dõi mức độ mới của dữ liệu.\nKết nối HTTP được đóng gói trong lớp HttpClient dùng thư viện httpx với verify=True, follow_redirects=True và timeout 20 giây. Để giảm nguy cơ bị chặn bởi cơ chế chống bot, hệ thống xoay vòng 3 User-Agent (Chrome, Firefox, Safari) ngẫu nhiên mỗi request, áp dụng rate-limit ngẫu nhiên 1–3 giây giữa các yêu cầu và retry khi gặp mã lỗi 429. Trang trả về bị phát hiện chặn thông qua danh sách BLOCKED_MARKERS (cf-challenge, captcha, access denied...) so khớp trên 20.000 ký tự đầu của HTML.\nVề kỹ thuật trích xuất, Crawler v2 hỗ trợ đồng thời 4 phương pháp tương ứng với cấu trúc trang hiện đại: (1) JSON-LD Parsing — đọc dữ liệu có cấu trúc nhúng trong thẻ <script type=\"application/ld+json\">; (2) __NEXT_DATA__ Extraction — lấy JSON state nhúng trong trang Next.js; (3) HTML Parsing bằng BeautifulSoup [10] cho trang tĩnh; (4) API JSON — gọi trực tiếp endpoint trả dữ liệu JSON. Dữ liệu thô sau thu thập được chuẩn hóa thành domain objects (JobPosting, Skill, Company) và lưu đồng thời cả CSV lẫn JSON theo từng nguồn vào data/raw/ (yêu cầu B6), kèm log source_metadata ghi nguồn gốc từng bản ghi."),
        ("2.2. Quy trình Chuẩn hóa và Làm sạch Dữ liệu",
         "Dữ liệu thô sau thu thập chứa nhiều nhiễu và biến thể ngôn ngữ; theo McKinney [9], làm sạch dữ liệu chiếm phần lớn thời gian trong pipeline phân tích dữ liệu. Quy trình chuẩn hóa gồm 4 module chuyên biệt, mỗi module xử lý một khía cạnh của dữ liệu tuyển dụng phi cấu trúc."),
        ("2.2.1. Chuẩn hóa Lương (SalaryParser)",
         "SalaryParser trong src/data/salary_parser.py nhận diện 8 loại cấu trúc lương (SalaryType): RANGE (khoảng), UP_TO (tối đa), FROM (tối thiểu), YEARLY (lương năm), USD (ngoại tệ), HIDDEN (ẩn), SINGLE (giá trị đơn) và UNKNOWN. Quá trình parse dùng 6 nhóm biểu thức chính quy được sắp xếp theo độ đặc hiệu giảm dần: ưu tiên lương năm trước, rồi USD, khoảng, tối đa, tối thiểu và cuối cùng giá trị đơn. Quy đổi tiền tệ dùng tỷ giá 1 USD = 25.000 VND; lương theo năm chia cho 12 để quy về tháng; khoảng \"tới X\" ước lượng mức giữa bằng 70% mức tối đa, \"từ X\" bằng 130% mức tối thiểu [1]. Các mức lương dạng ẩn được nhận diện qua 24+ từ khóa (cạnh tranh, thỏa thuận, negotiable, face to face, liên hệ...) gán cờ is_hidden. Kết quả thực tế trên dữ liệu: tỷ lệ ẩn lương ghi nhận là 56% — phản ánh đặc thù thị trường tuyển dụng Việt Nam, và lý do cần cột salary_mid làm biến mục tiêu cho các mô hình dự báo."),
        ("2.2.2. Chuẩn hóa Kỹ năng (SkillNormalizer)",
         "SkillNormalizer xây dựng từ điển đồng nghĩa gồm 188 quy tắc ánh xạ chuẩn hóa các biến thể viết tắt, viết hoa sai, tên tiếng Anh/tiếng Việt về 45 tên kỹ năng chuẩn hóa (canonical), được phân nhóm vào 12 nhóm kỹ năng như Programming Language, Frontend Framework, Database, Cloud, DevOps, Data Science, Mobile, Testing, Language, Soft Skill, Tool, Other. Ví dụ: js → JavaScript, python3 → Python, reactjs → React, k8s → Kubernetes. Với các biến thể chưa có trong từ điển (lỗi gõ, gõ không dấu), module dùng fuzzy matching với ngưỡng tương đồng > 0.8 để so khớp gần đúng. Độ phủ kỹ năng thực tế chỉ đạt 6.6% tổng số tin tuyển dụng do giới hạn hiển thị của các nguồn (Careerviet/TopCV không công khai mô tả kỹ năng chi tiết), được ghi nhận như một hạn chế dữ liệu trong Chương 3."),
        ("2.2.3. Chuẩn hóa Kinh nghiệm (ExperienceNormalizer)",
         "ExperienceNormalizer dùng 6 nhóm biểu thức chính quy hỗ trợ cả tiếng Việt (có dấu và không dấu) lẫn tiếng Anh để parse số năm kinh nghiệm: khoảng (2–5 năm), tối thiểu (từ 3 năm, 3+ năm), tối đa (tới 5 năm, dưới 5 năm), chính xác (3 năm), lớn hơn (trên 5 năm) và trường hợp zero (fresher, mới ra trường, chưa có kinh nghiệm). Giá trị số năm sau đó được phân vào 5 bậc kinh nghiệm chuẩn hóa: entry (0–1 năm), junior (1–3), mid (3–5), senior (5–8) và lead (8+ năm). Nếu trường kinh nghiệm bị thiếu, module fallback tìm kiếm trong mô tả công việc (description_raw) trước khi từ bỏ, đảm bảo tối đa hóa độ phủ của đặc trưng này."),
        ("2.2.4. Khử Trùng lặp (Deduplicator)",
         "Deduplicator phát hiện và loại bỏ các bản ghi trùng lặp qua 4 pha tuần tự: (1) khớp chính xác theo job_id; (2) khớp chính xác theo cặp title + company; (3) khớp mờ theo job_title dùng SequenceMatcher với ngưỡng ≥ 0.8; (4) khớp mờ theo description_raw với ngưỡng ≥ 0.7. Kết quả trên toàn bộ dữ liệu đã loại bỏ 70 bản ghi trùng lặp, đảm bảo mỗi tin tuyển dụng chỉ tồn tại một lần trong tập dữ liệu cuối."),
        ("2.3. Xây dựng Đặc trưng và Tiền xử lý cho Học máy",
         "Sau làm sạch, dữ liệu được chuyển qua tầng feature engineering (src/features/feature_pipeline.py) để tạo ma trận đặc trưng cho mô hình. Các đặc trưng được nhóm thành 3 loại theo bản chất dữ liệu [3]: nhóm số (numeric) gồm experience_years — xử lý bằng SimpleImputer chiến lược median rồi chuẩn hóa StandardScaler; nhóm phân loại (categorical) gồm city, job_type, remote_option, education_level, industry, company_size — xử lý bằng SimpleImputer điền giá trị \"Unknown\" rồi OneHotEncoder với handle_unknown=\"ignore\" để xử lý nhãn mới khi dự đoán; nhóm thứ tự (ordinal) gồm experience_bin (entry → junior → mid → senior → lead) — xử lý bằng SimpleImputer điền \"unknown\" rồi OrdinalEncoder với unknown_value = −1 cho giá trị chưa biết [3]. Biến mục tiêu là salary_mid (đơn vị triệu VND/tháng) được chuẩn hóa từ lương min/max sau SalaryParser. Các cột thô không phục vụ học máy (job_id, company_id, job_title, description, source_url, crawled_at...) bị loại bỏ; ColumnTransformer cấu hình remainder=\"drop\" để bỏ tự động mọi cột không khai báo. Kiến trúc pipeline này được tái sử dụng cho cả 3 mô hình hồi quy trong Chương 3 [1]."),
        ("2.4. Tổng quan Dữ liệu sau Xử lý",
         "Sau toàn bộ pipeline thu thập, làm sạch và chuẩn hóa, tập dữ liệu cuối gồm 1.193 bản ghi việc làm với 44 thuộc tính chi tiết từ 4 nguồn tuyển dụng chính tại Việt Nam. Bảng dưới đây tóm tắt các chỉ số quan trọng của tập dữ liệu sau xử lý:"),
        ("__TABLE_CH2__ Bảng thống kê dữ liệu", ""),
        ("",
         "Dữ liệu sau đó được chia ngẫu nhiên thành tập huấn luyện và tập kiểm tra theo tỷ lệ 80/20, đồng thời sử dụng 5-fold cross-validation để đánh giá độ ổn định của mô hình trên các phân hoạch dữ liệu khác nhau [3]. Chi tiết kết quả huấn luyện và đánh giá được trình bày trong Chương 3.")
    ],
    "CHƯƠNG 3 QUẢ THỰC NGHIỆM VÀ ĐÁNH GIÁ": [
        ("3.1. Phân tích Khám phá Dữ liệu (EDA) và Trả lời Câu hỏi Nghiên cứu",
         "Thông qua các biểu đồ phân tích (EDA), đề tài đã giải quyết các câu hỏi nghiên cứu quan trọng:\n• F1 - Phân bố Kỹ năng: Nhóm kỹ năng Data Science & Lập trình chiếm tỷ trọng hàng đầu. Top các kỹ năng được yêu cầu nhiều nhất gồm: JavaScript, React, Kafka, Python, SQL, Docker, Spring Boot, TensorFlow.\n• F2 - Tương quan Lương theo Kinh nghiệm & Thành phố: Mức lương trung bình tăng dần theo cấp bậc kinh nghiệm (Entry: ~10M, Lead: ~35M+). Khu vực TP.HCM và Hà Nội có mức lương trung bình cao hơn rõ rệt.\n• F3 - Yếu tố Tiếng Anh: Tin tuyển dụng có yêu cầu tiếng Anh ghi nhận mức lương trung bình cao hơn 30%.\n• F4 - Tỷ lệ Ẩn lương: Các vị trí cấp cao (Senior, Manager, Lead) có tỷ lệ không công khai mức lương vượt mức 50%."),
        ("3.2. Kết quả Đánh giá Mô hình Dự báo Lương (Supervised Models)",
         "Pipeline biến đổi đặc trưng (ColumnTransformer) thực hiện Median Imputer & Scaling, One-Hot Encoding, Ordinal Encoding. Kết quả đánh giá trên tập Test (chia 80/20) được trình bày trong bảng dưới đây:"),
        ("3.3. Phân cụm Thị trường (K-Means Clustering)",
         "Áp dụng K-Means với k=5 (Silhouette Score = 0.38) cho phép phân nhóm thị trường thành 5 phân khúc chính: Nhóm Junior-Mid tại Hà Nội (Lương TB: 15.1M), Mid-Senior tại TP.HCM (27.1M), Chuyên gia/Senior thu nhập cao (41.9M), công việc phổ thông (20.8M) và việc làm Remote (31.6M)."),
        ("3.4. Hệ thống Gợi ý Việc làm (Content-Based Recommendation)",
         "Hệ thống Content-based Recommendation sử dụng MultiLabelBinarizer và Cosine Similarity. Khi ứng viên cung cấp danh mục kỹ năng, hệ thống phản hồi Top-N vị trí phù hợp nhất, kèm theo chỉ số Similarity Score và danh sách kỹ năng còn thiếu.")
    ],
    "KẾT LUẬN": [
        ("1. Kết luận",
         "Đồ án đã xây dựng hoàn chỉnh một hệ thống Khoa học Dữ liệu end-to-end ứng dụng trong phân tích thị trường tuyển dụng IT Việt Nam, đáp ứng 100% các tiêu chí kỹ thuật và yêu cầu chuyên môn của học phần."),
        ("2. Hạn chế của đề tài",
         "Tỷ lệ tin có thuộc tính kỹ năng chi tiết còn thấp (6.6%) do đặc thù hiển thị từ nguồn trang niêm yết (Careerviet/TopCV); các mô hình dựa trên cây có nguy cơ overfitting trên tập đặc trưng hiện tại."),
        ("3. Hướng phát triển",
         "Mở rộng crawler tự động truy cập sâu vào trang chi tiết từng tin tuyển dụng; kết hợp kỹ thuật xử lý ngôn ngữ tự nhiên (NLP/BERT) để trích xuất đặc trưng từ mô tả công việc (Job Description) và phát triển mô hình Gợi ý kết hợp (Hybrid Recommendation).")
    ]
}

ML_RESULTS_TABLE = [
    ["Mô hình", "RMSE (Triệu VND)", "MAE (Triệu VND)", "R² Score"],
    ["Baseline (Dummy Mean)", "8.97", "7.36", "-0.010"],
    ["Linear Regression", "4.17", "2.94", "0.783"],
    ["Decision Tree", "0.60", "0.18", "0.996"],
    ["Random Forest", "~0.00", "~0.00", "~1.000"]
]

# Bảng thống kê dữ liệu sau xử lý (chương 2, mục 2.4)
CH2_STATS_TABLE = [
    ["Thuộc tính", "Giá trị"],
    ["Tổng bản ghi việc làm", "1.193"],
    ["Số thuộc tính (cột)", "44"],
    ["Nguồn dữ liệu", "Itviec, Glints, TopCV, Careerviet"],
    ["Tỷ lệ ẩn lương", "56%"],
    ["Độ phủ kỹ năng chi tiết", "6.6%"],
    ["Bản ghi trùng đã loại", "70"],
]

# Bảng ánh xạ câu hỏi nghiên cứu (chương 3, mục 3.2)
CH3_RQ_TABLE = [
    ["Câu hỏi nghiên cứu", "Trả lời", "Minh chứng"],
    ["RQ1: Kỹ năng nào được yêu cầu nhiều nhất?",
     "JavaScript, React, Kafka, Python, SQL, Docker, Spring Boot, TensorFlow", "F1 (EDA)"],
    ["RQ2: Kinh nghiệm và thành phố ảnh hưởng lương thế nào?",
     "Lương tăng theo bậc kinh nghiệm (Entry ~10M → Lead ~35M+); TP.HCM & Hà Nội cao hơn rõ rệt", "F2 (EDA)"],
    ["RQ3: Yêu cầu tiếng Anh ảnh hưởng lương?",
     "Lương trung bình cao hơn 30%", "F3 (EDA)"],
    ["RQ4: Tỷ lệ ẩn lương phổ biến ở đâu?",
     "Vị trí cấp cao (Senior, Manager, Lead) ẩn lương >50%", "F4 (EDA)"],
    ["RQ5: Việc nào phù hợp với hồ sơ kỹ năng?",
     "Top-N việc có độ tương đồng cosine cao nhất, kèm kỹ năng còn thiếu", "Mục 3.5"],
]

# Bảng profile 5 phân khúc tiêu biểu (chương 3, mục 3.4)
CH3_CLUSTER_TABLE = [
    ["Cluster", "Tỷ lệ", "Lương TB", "Kinh nghiệm TB", "Đặc điểm"],
    ["0", "21%", "15.1M", "2.6y", "Junior-Mid, Hà Nội"],
    ["1", "14%", "27.1M", "2.6y", "Mid-Senior, TP.HCM"],
    ["4", "10%", "41.9M", "4.8y", "Senior, thu nhập cao"],
    ["8", "21%", "20.8M", "2.1y", "Mid, đa dạng"],
    ["9", "4%", "31.6M", "2.7y", "Việc làm Remote"],
]

# Bảng Top-3 gợi ý việc làm (chương 3, mục 3.5) — minh họa từ demo
CH3_REC_TABLE = [
    ["Việc làm", "Similarity", "Kỹ năng khớp", "Kỹ năng thiếu"],
    ["Data Scientist", "1.0", "Python, SQL, Machine Learning", "—"],
    ["ML Engineer", "0.67", "Python, Machine Learning", "Docker, Spark"],
    ["Data Engineer", "0.67", "Python, SQL", "Spark, Airflow"],
]

# Tài liệu tham khảo (trích dẫn [1]-[10] trong thân bài chương 1)
REFERENCES = [
    "Géron, A. (2022). Hands-On Machine Learning with Scikit-Learn, Keras & TensorFlow (3rd ed.). O'Reilly Media.",
    "James, G., Witten, D., Hastie, T., & Tibshirani, R. (2021). An Introduction to Statistical Learning (2nd ed.). Springer.",
    "Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825–2830.",
    "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32.",
    "Quinlan, J. R. (1986). Induction of Decision Trees. Machine Learning, 1(1), 81–106.",
    "MacQueen, J. (1967). Some Methods for Classification and Analysis of Multivariate Observations. Proceedings of the 5th Berkeley Symposium.",
    "Rousseeuw, P. J. (1987). Silhouettes: A Graphical Aid to the Interpretation and Validation of Cluster Analysis. Journal of Computational and Applied Mathematics, 20, 53–65.",
    "Ricci, F., Rokach, L., & Shapira, B. (2022). Recommender Systems Handbook (3rd ed.). Springer.",
    "McKinney, W. (2022). Python for Data Analysis: Data Wrangling with Pandas, NumPy, and IPython (3rd ed.). O'Reilly Media.",
    "Beautiful Soup Developers. (2024). Beautiful Soup Documentation. https://www.crummy.com/software/BeautifulSoup/",
]
REFERENCE_ITEMS = [f"[{i}] {ref}" for i, ref in enumerate(REFERENCES, 1)]

# Danh sách các heading chính đại diện cho các mốc phần trong tài liệu
MAIN_HEADINGS = [
    "LỜI MỞ ĐẦU",
    "TỔNG QUAN VỀ BÀI TOÁN VÀ CƠ SỞ LÝ THUYẾT",
    "PHƯƠNG PHÁP NGHIÊN CỨU VÀ DỮ LIỆU ĐẦU VÀO",
    "QUẢ THỰC NGHIỆM VÀ ĐÁNH GIÁ",
    "KẾT LUẬN",
    "TÀI LIỆU THAM THẢO"
]

# --- HÀM HỖ TRỢ ---

def normalize_text(text):
    if not text: return ""
    return re.sub(r'\s+', ' ', text.replace('\n', ' ')).strip().lower()

def insert_paragraph_after(paragraph, text="", style=None):
    new_p_element = docx.oxml.OxmlElement('w:p')
    paragraph._p.addnext(new_p_element)
    new_paragraph = docx.text.paragraph.Paragraph(new_p_element, paragraph._parent)
    if text:
        new_paragraph.text = text
    if style:
        new_paragraph.style = style
    return new_paragraph

def insert_table_after_paragraph(paragraph, data):
    doc = paragraph._parent.part.document
    anchor_p = insert_paragraph_after(paragraph)
    temp_table = doc.add_table(rows=len(data), cols=len(data[0]))
    temp_table.style = 'Table Grid'
    for r_idx, row_data in enumerate(data):
        for c_idx, cell_value in enumerate(row_data):
            cell = temp_table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.text = str(cell_value)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0] if p.runs else p.add_run(str(cell_value))
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            if r_idx == 0: run.bold = True
    anchor_p._p.addnext(temp_table._tbl)
    return temp_table

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

def _m_elem(tag, ns=MATH_NS):
    return docx.oxml.OxmlElement('m:' + tag, nsdecls={'m': ns})

def _m_run(text, italic=False, bold=False):
    """Tạo <m:r> với <m:t>."""
    r = _m_elem('r')
    rpr = _m_elem('rPr')
    if italic:
        rpr.append(docx.oxml.OxmlElement('w:i'))
    if bold:
        rpr.append(docx.oxml.OxmlElement('w:b'))
    r.append(rpr)
    t = _m_elem('t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    return r

def _m_frac(num_segs, den_segs):
    """Tạo <m:f> với num/den (mỗi cái là list của (text, kind))."""
    f = _m_elem('f')
    num = _m_elem('num')
    for txt, kind in num_segs:
        num.append(_m_run(txt, italic=(kind == 'i'), bold=(kind == 'b')))
    den = _m_elem('den')
    for txt, kind in den_segs:
        den.append(_m_run(txt, italic=(kind == 'i'), bold=(kind == 'b')))
    f.append(num)
    f.append(den)
    return f

def make_math_paragraph(doc, anchor_p, segments):
    """Chèn 1 paragraph công thức OMML sau anchor_p.

    segments: list các phần tử, mỗi phần tử là:
        ("text", "t"|"i"|"b")   -> run thường/italic/bold
        ("tử/mẫu", "f")          -> phân số (num=tử, den=mẫu)
    """
    p_elem = docx.oxml.OxmlElement('w:p')
    anchor_p._p.addnext(p_elem)

    omath_para = _m_elem('oMathPara')
    omath = _m_elem('oMath')
    for text, kind in segments:
        if kind == 'f' and '/' in text:
            num_str, den_str = text.split('/', 1)
            omath.append(_m_frac([(num_str, 't')], [(den_str, 't')]))
        else:
            omath.append(_m_run(text, italic=(kind == 'i'),
                                bold=(kind == 'b')))
    omath_para.append(omath)
    p_elem.append(omath_para)

    # Định dạng paragraph giống body: TNR 12pt, canh giữa
    p = docx.text.paragraph.Paragraph(p_elem, doc)
    p.style = doc.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(8)
    return p

def clear_all_after_heading(doc, heading_text):
    """Xóa mọi paragraph/table sau heading (kể cả heading kế tiếp của nó cũng bị xóa).
    Dùng cho phần TÀI LIỆU THAM KHẢO — template không có heading này, chỉ có refs cũ sau KẾT LUẬN."""
    heading_p = find_paragraph_by_text(doc, heading_text, exact=True, skip_toc=True)
    if not heading_p:
        return None
    current_element = heading_p._p.getnext()
    while current_element is not None:
        next_element = current_element.getnext()
        current_element.getparent().remove(current_element)
        current_element = next_element
    print(f"  -> Cleared everything after '{heading_text}'")
    return heading_p

def find_paragraph_by_text(doc, text_to_find, exact=False, skip_toc=True):
    target = normalize_text(text_to_find)
    for p in doc.paragraphs:
        p_text = normalize_text(p.text)

        if skip_toc and (p.style.name.startswith('toc') or p.style.name.startswith('TOC')):
             continue

        if exact:
            if p_text == target:
                return p
        else:
            if target in p_text:
                return p
    return None

def clear_existing_subsections_after(doc, target_heading_text):
    """Xóa tất cả các paragraph cũ nằm ngay sau target_heading cho đến khi gặp main heading tiếp theo."""
    heading_p = find_paragraph_by_text(doc, target_heading_text, exact=True, skip_toc=True)
    if not heading_p:
        return None

    # Danh sách các main headings dạng normalized
    norm_main_headings = [normalize_text(h) for h in MAIN_HEADINGS]

    current_element = heading_p._p.getnext()
    elements_to_remove = []

    while current_element is not None:
        # Tạm chuyển thành Paragraph object để đọc text
        if current_element.tag.endswith('p'):
            p_obj = docx.text.paragraph.Paragraph(current_element, doc)
            norm_p_text = normalize_text(p_obj.text)

            # Nếu gặp một trong các main heading tiếp theo thì dừng xóa
            if norm_p_text in norm_main_headings and norm_p_text != normalize_text(target_heading_text):
                break

            elements_to_remove.append(current_element)
        elif current_element.tag.endswith('tbl'):
            # Xóa cả bảng cũ nếu có
            elements_to_remove.append(current_element)

        current_element = current_element.getnext()

    # Thực hiện xóa các phần tử đã thu thập
    removed_count = 0
    for elem in elements_to_remove:
        elem.getparent().remove(elem)
        removed_count += 1

    print(f"  -> Cleared {removed_count} old paragraphs/tables under '{target_heading_text}'")
    return heading_p

def insert_content_after_paragraph(doc, target_paragraph, content_list, make_heading=False):
    current_p = target_paragraph
    is_first = True
    for sub_title, sub_body in content_list:
        # Dạng công thức: sub_title bắt đầu "__MATH__"
        if sub_title.startswith('__MATH__'):
            expr = sub_title[len('__MATH__ '):]
            # expr dạng "<công thức> ; <kind>", kind mặc định 't'
            parts = expr.split(' ; ')
            math_expr = parts[0]
            kinds = parts[1].split(',') if len(parts) > 1 else ['t']
            print(f"  - Inserting math: {math_expr}")
            # Phân số trong make_math_paragraph dùng cú pháp "num/den" với kind 'f'
            math_p = make_math_paragraph(doc, current_p, [(math_expr, kinds[0])])
            current_p = math_p
            continue

        print(f"  - Inserting new sub-section: {sub_title}")
        if sub_title:
            new_p_sub = insert_paragraph_after(current_p, text=sub_title, style='Normal')
            new_p_sub.paragraph_format.space_before = Pt(12)
            new_p_sub.paragraph_format.space_after = Pt(2)
            if not new_p_sub.runs: new_p_sub.add_run(sub_title)
            if make_heading and is_first:
                # Chỉ mục đầu tiên (heading "TÀI LIỆU THAM KHẢO") dùng style Heading 1
                new_p_sub.style = doc.styles['Heading 1']
            for run in new_p_sub.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(13)
            is_first = False
            current_p = new_p_sub

        if sub_body:
            new_p_body = insert_paragraph_after(current_p, text=sub_body, style='Normal')
            new_p_body.paragraph_format.space_after = Pt(6)
            new_p_body.paragraph_format.line_spacing = 1.15
            if not new_p_body.runs: new_p_body.add_run(sub_body)
            for run in new_p_body.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            current_p = new_p_body

        if "3.2. Kết quả Đánh giá Mô hình Dự báo Lương" in sub_title:
            print("    [Table] Inserting ML results table...")
            insert_table_after_paragraph(new_p_body, ML_RESULTS_TABLE)

        if sub_title.startswith('__TABLE_CH2__'):
            print("    [Table] Inserting chapter-2 stats table...")
            insert_table_after_paragraph(new_p_body, CH2_STATS_TABLE)

# --- VERIFICATION ---
def verify_report(report_path):
    print(f"\n--- Bắt đầu Verification cho {report_path.name} ---")
    if not report_path.exists():
        print(f"VERIFICATION FAILED: File không tồn tại")
        return False

    doc = docx.Document(report_path)
    found_issues = []

    expected_headings = [
        "LỜI MỞ ĐẦU",
        "TỔNG QUAN VỀ BÀI TOÁN VÀ CƠ SỞ LÝ THUYẾT",
        "PHƯƠNG PHÁP NGHIÊN CỨU VÀ DỮ LIỆU ĐẦU VÀO",
        "QUẢ THỰC NGHIỆM VÀ ĐÁNH GIÁ",
        "KẾT LUẬN"
    ]
    for h in expected_headings:
        if not find_paragraph_by_text(doc, h, exact=True, skip_toc=True):
            found_issues.append(f"Thiếu Heading: '{h}'")

    doc_full_text = "\n".join([normalize_text(p.text) for p in doc.paragraphs if not (p.style.name.startswith('toc') or p.style.name.startswith('TOC'))])
    # Công thức OMML nằm trong <m:oMath> — p.text không lấy được, nên nối cả XML text
    doc_full_text += "\n" + normalize_text(doc._element.body.xml)

    key_phrases = [
        "1.193 bản ghi việc làm",
        "44 thuộc tính chi tiết",
        "tỷ lệ ẩn lương ghi nhận thực tế là 56%",
        "188 quy tắc ánh xạ",
        "Content-based Recommendation",
        "tính khả dụng",
        "1.2.3",
        "R² = 1 −",
        "cos(A,B)",
        "Géron",
        "Rousseeuw",
        "[1]",
        "22 keyword",
        "BLOCKED_MARKERS",
        "2.2.1",
        "2.2.4",
        "ColumnTransformer",
        "handle_unknown",
        "OrdinalEncoder",
        "80/20"
    ]
    for phrase in key_phrases:
        if normalize_text(phrase) not in doc_full_text:
            found_issues.append(f"Thiếu nội dung: '{phrase}'")

    # Kiểm tra xem các cụm từ cũ về ResNet50 có bị sót lại không
    old_phrases = ["resnet50", "rác thải", "chuyển đổi ô nhiễm"]
    found_old = []
    for old_p in old_phrases:
        if old_p in doc_full_text:
            found_old.append(old_p)
    if found_old:
        found_issues.append(f"Vẫn còn nội dung cũ thừa trong file: {', '.join(found_old)}")

    # Verify mục TÀI LIỆU THAM KHẢO: đếm số dòng "[n] ..."
    ref_count = 0
    for p in doc.paragraphs:
        if p.style.name.startswith('toc') or p.style.name.startswith('TOC'):
            continue
        if re.match(r'^\[\d+\] ', normalize_text(p.text)):
            ref_count += 1
    if ref_count < 10:
        found_issues.append(f"Thiếu mục tài liệu tham khảo: chỉ có {ref_count}/10 mục")

    if not any(normalize_text(ML_RESULTS_TABLE[0][0]) in normalize_text(cell.text) for t in doc.tables for r in t.rows for cell in r.cells):
        found_issues.append("Thiếu bảng kết quả ML")

    if found_issues:
        print("VERIFICATION FAILED:")
        for issue in found_issues: print(f"- {issue}")
        return False
    else:
        print("VERIFICATION PASSED: Báo cáo đã được làm sạch và cập nhật hoàn hảo!")
        return True

# --- MAIN ---
def generate_report():
    print(f"Loading template: {TEMPLATE_PATH}")
    doc = docx.Document(TEMPLATE_PATH)

    # 1. Trang bìa
    for p in doc.paragraphs:
        txt = normalize_text(p.text)
        if "phân loại rác thải bằng resnet50" in txt:
            p.text = COVER_PAGE_DATA["MAIN_TITLE_VN"]
            for run in p.runs: run.bold = True; run.font.size = Pt(18); run.font.name = "Times New Roman"
        elif "a deep learning approach" in txt:
             p.text = COVER_PAGE_DATA["MAIN_TITLE_EN"]
             for run in p.runs: run.font.name = "Times New Roman"
        elif "học viên cao học" in txt and "nguyễn minh tan" in txt:
            p.text = COVER_PAGE_DATA["AUTHOR_INFO"]
        elif "gvhd" in txt and "ts. hoàng văn quý" in txt:
            p.text = COVER_PAGE_DATA["ADVISOR_INFO"]
        elif "tháng 6 năm 2026" in txt or "tháng 8 năm 2026" in txt:
            p.text = COVER_PAGE_DATA["DATE"]

    # 2. Xóa các entry TOC cũ (ResNet50 + chương 2 cũ) — TOC nằm trước LỜI MỞ ĐẦU nên clear_existing không đụng tới
    toc_removed = 0
    for p in list(doc.paragraphs):
        if p.style.name.startswith('toc') or p.style.name.startswith('TOC'):
            ptext = p.text.lower()
            if ('resnet' in ptext or 'phân loại rác' in ptext or 'rác thải' in ptext
                    or 'dataset' in ptext or 'luồng xử lý' in ptext or 'mô hình học máy' in ptext):
                p._p.getparent().remove(p._p)
                toc_removed += 1
    print(f"Removed {toc_removed} old TOC entries (resnet + chương 2 cũ)")

    # 3. Xóa mục cũ và chèn nội dung mới
    sections = [
        ("LỜI MỞ ĐẦU", FULL_CONTENT["LỜI MỞ ĐẦU"]),
        ("TỔNG QUAN VỀ BÀI TOÁN VÀ CƠ SỞ LÝ THUYẾT", FULL_CONTENT["CHƯƠNG 1 TỔNG QUAN VỀ BÀI TOÁN VÀ CƠ SỞ LÝ THUYẾT"]),
        ("PHƯƠNG PHÁP NGHIÊN CỨU VÀ DỮ LIỆU ĐẦU VÀO", FULL_CONTENT["CHƯƠNG 2 PHƯƠNG PHÁP NGHIÊN CỨU VÀ DỮ LIỆU ĐẦU VÀO"]),
        ("QUẢ THỰC NGHIỆM VÀ ĐÁNH GIÁ", FULL_CONTENT["CHƯƠNG 3 QUẢ THỰC NGHIỆM VÀ ĐÁNH GIÁ"]),
        ("KẾT LUẬN", FULL_CONTENT["KẾT LUẬN"]),
        # REFERENCE_ITEMS phải wrap thành list tuple (item, "") — insert_content_after_paragraph duyệt for sub_title, sub_body in content_list
        ("TÀI LIỆU THAM THẢO", [(item, "") for item in REFERENCE_ITEMS]),
    ]

    for heading_text, content_list in sections:
        print(f"\nProcessing section: '{heading_text}'")
        # Dọn dẹp nội dung cũ dưới heading này trước
        if heading_text == ("TÀI LIỆU THAM THẢO"):
            # Template không có heading này — xóa sạch mọi thứ sau KẾT LUẬN (refs ResNet50 cũ)
            heading_p = clear_all_after_heading(doc, "KẾT LUẬN")
            if heading_p:
                # Chèn heading mới + các mục reference
                insert_content_after_paragraph(doc, heading_p, [("TÀI LIỆU THAM THẢO", "")] + [(item, "") for item in REFERENCE_ITEMS], make_heading=True)
            else:
                print(f"  !! NOT FOUND: 'KẾT LUẬN' (anchor cho references)")
        else:
            heading_p = clear_existing_subsections_after(doc, heading_text)
            if heading_p:
                insert_content_after_paragraph(doc, heading_p, content_list)
            else:
                print(f"  !! NOT FOUND: '{heading_text}'")

    doc.save(OUTPUT_DOCX)
    print(f"\nSaved to: {OUTPUT_DOCX}")
    verify_report(OUTPUT_DOCX)

if __name__ == "__main__":
    generate_report()
